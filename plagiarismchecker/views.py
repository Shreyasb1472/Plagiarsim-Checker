from django.shortcuts import render
from django.http import HttpResponse
from plagiarismchecker.algorithm import main
from plagiarismchecker.algorithm import fileSimilarity
from docx import Document
from PyPDF2 import PdfReader

# Home page
def home(request):
    return render(request, 'pc/index.html')

# Web search using text input
def test(request):
    print("request is welcome test")
    print(request.POST['q'])  

    if request.POST['q']: 
        percent, link = main.findSimilarity(request.POST['q'])
        percent = round(percent, 2)
    else:
        percent, link = 0.0, ""

    print("Output.....................!!!!!!!!", percent, link)
    return render(request, 'pc/index.html', {'link': link, 'percent': percent})

# Web search using uploaded file (.txt, .docx, .pdf)
def filetest(request):
    value = ''    
    uploaded_file = request.FILES.get('docfile')

    if not uploaded_file:
        return render(request, 'pc/index.html', {'link': '', 'percent': 0.0, 'error': 'No file uploaded'})

    filename = str(uploaded_file.name).lower()

    if filename.endswith(".txt"):
        value = uploaded_file.read().decode('utf-8', errors='ignore')

    elif filename.endswith(".docx"):
        document = Document(uploaded_file)
        for para in document.paragraphs:
            value += para.text

    elif filename.endswith(".pdf"):
        reader = PdfReader(uploaded_file)
        for page in reader.pages:
            text = page.extract_text()
            if text:
                value += text

    if not value.strip():
        return render(request, 'pc/index.html', {'link': '', 'percent': 0.0, 'error': 'Uploaded file is empty or unreadable.'})

    percent, link = main.findSimilarity(value)
    percent = round(percent, 2)
    print("Output...................!!!!!!!!", percent, link)
    return render(request, 'pc/index.html', {'link': link, 'percent': percent})

# Show the compare documents page
def fileCompare(request):
    return render(request, 'pc/doc_compare.html')

# Compare two raw text inputs
def twofiletest1(request):
    print("Submitted text for 1st and 2nd")
    text1 = request.POST.get('q1', '')
    text2 = request.POST.get('q2', '')

    if text1.strip() and text2.strip():
        result = fileSimilarity.findFileSimilarity(text1, text2)
        result = round(result, 2)
    else:
        result = "Error: One or both input fields are empty."

    print("Output>>>>>>>>>>>>>>>>>>>>!!!!!!!!", result)
    return render(request, 'pc/doc_compare.html', {'result': result})

# Compare two uploaded files (.txt, .docx, .pdf)
def twofilecompare1(request):
    value1 = ''
    value2 = ''
    file1 = request.FILES.get('docfile1')
    file2 = request.FILES.get('docfile2')

    if not file1 or not file2:
        return render(request, 'pc/doc_compare.html', {'result': "Error: Both files must be uploaded."})

    filename1 = file1.name.lower()
    filename2 = file2.name.lower()

    try:
        if filename1.endswith(".txt"):
            value1 = file1.read().decode('utf-8', errors='ignore')
        elif filename1.endswith(".docx"):
            doc1 = Document(file1)
            for para in doc1.paragraphs:
                value1 += para.text
        elif filename1.endswith(".pdf"):
            reader1 = PdfReader(file1)
            for page in reader1.pages:
                text = page.extract_text()
                if text:
                    value1 += text

        if filename2.endswith(".txt"):
            value2 = file2.read().decode('utf-8', errors='ignore')
        elif filename2.endswith(".docx"):
            doc2 = Document(file2)
            for para in doc2.paragraphs:
                value2 += para.text
        elif filename2.endswith(".pdf"):
            reader2 = PdfReader(file2)
            for page in reader2.pages:
                text = page.extract_text()
                if text:
                    value2 += text
    except Exception as e:
        return render(request, 'pc/doc_compare.html', {'result': f"Error processing files: {str(e)}"})

    if not value1.strip() or not value2.strip():
        return render(request, 'pc/doc_compare.html', {'result': "Error: One or both uploaded files are empty or unreadable."})

    result = fileSimilarity.findFileSimilarity(value1, value2)
    result = round(result, 2)
    print("Output..................!!!!!!!!", result)

    return render(request, 'pc/doc_compare.html', {'result': result})
